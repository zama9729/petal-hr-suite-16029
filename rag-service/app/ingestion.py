"""Document ingestion pipeline."""
import hashlib
import os
from typing import List, Dict, Optional
from pathlib import Path
import pypdf
import pdfplumber
from docx import Document as DocxDocument
from app.models import Document, DocumentChunk, Tenant
from app.pii_detection import pii_detector
from app.config import settings
from sqlalchemy.orm import Session
import logging
import uuid

logger = logging.getLogger(__name__)


class DocumentIngester:
    """Document ingestion and chunking."""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from document."""
        try:
            if file_type == "pdf":
                return self._extract_from_pdf(file_path)
            elif file_type == "docx":
                return self._extract_from_docx(file_path)
            elif file_type in ["txt", "md"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            # Fallback: try OCR (commented for now, requires tesseract setup)
            # return self._extract_with_ocr(file_path)
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber (better than pypdf for complex PDFs)."""
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying pypdf: {e}")
            # Fallback to pypdf
            with open(file_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        return "\n\n".join([para.text for para in doc.paragraphs])
    
    def chunk_text(self, text: str, metadata: Optional[Dict] = None) -> List[Dict]:
        """Chunk text with overlap."""
        # Simple token-based chunking (in production, use tiktoken for accurate token counting)
        words = text.split()
        chunks = []
        chunk_index = 0
        
        i = 0
        while i < len(words):
            # Approximate chunk size (roughly 4 chars per token)
            chunk_words = []
            chunk_char_count = 0
            
            while i < len(words) and chunk_char_count < self.chunk_size * 4:
                chunk_words.append(words[i])
                chunk_char_count += len(words[i]) + 1  # +1 for space
                i += 1
            
            chunk_text = " ".join(chunk_words)
            
            if chunk_text.strip():
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": chunk_index,
                    "metadata": metadata or {}
                })
                chunk_index += 1
            
            # Overlap: move back by overlap tokens
            if i < len(words):
                overlap_tokens = int(self.chunk_overlap * 4 / 5)  # Rough estimate
                i = max(0, i - overlap_tokens)
        
        return chunks
    
    def process_document(
        self,
        file_path: str,
        tenant_id: uuid.UUID,
        filename: str,
        is_confidential: bool = False,
        db: Optional[Session] = None
    ) -> Document:
        """Process and ingest document."""
        file_type = Path(filename).suffix[1:].lower()
        
        if file_type not in settings.allowed_extensions:
            raise ValueError(f"File type {file_type} not allowed")
        
        # Extract text
        text = self.extract_text(file_path, file_type)
        
        # Calculate content hash
        content_hash = hashlib.sha256(text.encode()).hexdigest()
        
        # Check for duplicate
        if db:
            existing = db.query(Document).filter(
                Document.tenant_id == tenant_id,
                Document.content_hash == content_hash
            ).first()
            if existing:
                logger.info(f"Duplicate document found: {existing.id}")
                return existing
        
        # Create document record
        doc = Document(
            tenant_id=tenant_id,
            filename=filename,
            file_path=file_path,
            file_type=file_type,
            file_size=os.path.getsize(file_path),
            content_hash=content_hash,
            is_confidential=is_confidential,
            ingestion_status="processing",
            meta_data={"filename": filename, "file_type": file_type}
        )
        
        if db:
            db.add(doc)
            db.commit()
            db.refresh(doc)
        
        # Chunk text
        chunks = self.chunk_text(text, metadata={"filename": filename, "file_type": file_type})
        
        # Create chunk records
        chunk_records = []
        for chunk_data in chunks:
            # Redact PII from chunk content
            content_redacted = pii_detector.redact_pii(chunk_data["text"])
            
            chunk = DocumentChunk(
                document_id=doc.id,
                tenant_id=tenant_id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["text"],
                content_redacted=content_redacted,
                chunk_metadata=chunk_data.get("metadata", {})
            )
            chunk_records.append(chunk)
        
        if db:
            db.add_all(chunk_records)
            db.commit()
        
        doc.ingestion_status = "completed"
        if db:
            db.commit()
        
        logger.info(f"Processed document {doc.id} with {len(chunk_records)} chunks")
        return doc


# Global instance
document_ingester = DocumentIngester()

